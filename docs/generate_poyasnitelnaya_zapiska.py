# -*- coding: utf-8 -*-
"""Генерация пояснительной записки в формате Microsoft Word (.docx)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

FONT_NAME = "Times New Roman"
RGB_BLACK = RGBColor(0, 0, 0)


def _set_run_font(
    run,
    *,
    size_pt: int | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
) -> None:
    run.font.name = FONT_NAME
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def _style_heading_runs(paragraph, size_pt: int) -> None:
    for run in paragraph.runs:
        _set_run_font(run, size_pt=size_pt, color=RGB_BLACK)


def _add_hyperlink(paragraph: Paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    r_pr.append(color)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    r_pr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(14 * 2)))
    r_pr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(int(14 * 2)))
    r_pr.append(sz_cs)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _p(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size_pt=14, color=RGB_BLACK)
    if bold:
        run.bold = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(1.25)


def _bullets(doc: Document, items: list[str]) -> None:
    for t in items:
        p = doc.add_paragraph(t, style="List Bullet")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            _set_run_font(run, size_pt=14, color=RGB_BLACK)


def _h2(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=2)
    _style_heading_runs(h, 14)


def _set_table_body_font(tbl, size_pt: int = 14) -> None:
    for row in tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    _set_run_font(run, size_pt=size_pt, color=RGB_BLACK)


def _apply_page_margins(doc: Document) -> None:
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)


def build() -> Document:
    doc = Document()
    _apply_page_margins(doc)
    for _sn in ("Normal", "Heading 1", "Heading 2", "List Paragraph", "List Number", "List Bullet"):
        try:
            doc.styles[_sn].font.name = FONT_NAME
        except KeyError:
            pass

    t = doc.add_paragraph()
    t.alignment = 1  # CENTER
    r = t.add_run("ПОЯСНИТЕЛЬНАЯ ЗАПИСКА")
    _set_run_font(r, size_pt=16, color=RGB_BLACK)
    r.bold = True
    t = doc.add_paragraph()
    t.alignment = 1
    r2 = t.add_run(
        "к учебному проекту «Список задач (TODO)» — fullstack-приложение\n"
        "(backend: Python / FastAPI; веб-клиент; мобильный клиент: Flutter)"
    )
    _set_run_font(r2, size_pt=14, color=RGB_BLACK)

    doc.add_paragraph()

    h = doc.add_heading("1. Описание предметной области", level=1)
    _style_heading_runs(h, 16)
    _p(
        doc,
        "Предметная область — персональное планирование и учёт бытовых и учебных дел в виде "
        "электронного списка задач. Пользователь работает в личном кабинете: создаёт записи "
        "о делах (заголовок, при необходимости — описание), отмечает выполнение, редактирует "
        "и удаляет позиции. Каждая учётная запись изолирована: задачи одного пользователя "
        "недоступны другому. Для демонстрации разграничения прав в учебных целях введена роль "
        "администратора с отдельным сценарием просмотра (все задачи в системе).",
    )

    h = doc.add_heading("2. Цели и задачи", level=1)
    _style_heading_runs(h, 16)
    _p(
        doc,
        "Цель проекта — спроектировать и реализовать полный контур fullstack-приложения: "
        "реляционное хранение данных, REST API с аутентификацией и авторизацией, веб-интерфейс "
        "и кроссплатформенный мобильный клиент, согласованные по одному контракту API.",
    )
    _p(doc, "Задачи:", bold=True)
    _bullets(
        doc,
        [
            "Спроектировать модель «пользователь — задача» и отразить её в СУБД с помощью ORM и миграций.",
            "Реализовать для сущности «задача» полный набор операций CRUD с проверкой прав доступа по владельцу.",
            "Обеспечить регистрацию и вход пользователя, выдачу и проверку JWT при обращении к защищённым маршрутам.",
            "Реализовать веб-клиент (HTML/CSS/JavaScript) и клиент на Flutter с сохранением токена и вызовами REST.",
            "Задокументировать стек технологий, инструменты статического анализа, СУБД, логирование, тестирование и подходы к контролю доступа.",
        ],
    )

    h = doc.add_heading("3. Функциональные требования", level=1)
    _style_heading_runs(h, 16)
    _p(doc, "К системе предъявлены следующие функциональные требования.")
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Функция"
    hdr[1].text = "Описание"
    rows = [
        ("Регистрация", "Создание учётной записи по email и паролю."),
        ("Вход", "Аутентификация, получение JWT (OAuth2 password flow; в форме поле username соответствует email)."),
        ("Список задач", "GET /tasks — выдача только задач текущего пользователя; поддержка фильтра по признаку выполнения и сортировки."),
        ("Создание задачи", "POST /tasks — новая задача привязывается к идентификатору владельца."),
        ("Просмотр одной задачи", "GET /tasks/{id} — доступ только при совпадении владельца."),
        ("Обновление", "PATCH /tasks/{id} — частичное изменение полей (заголовок, признак выполнения и др.)."),
        ("Удаление", "DELETE /tasks/{id} — только для владельца записи; успех 204 без тела."),
        ("Профиль", "GET /auth/me — id, email, роль из БД; клиенты используют для переключения на /admin/tasks."),
        ("Админ: список", "GET /admin/tasks — все задачи; фильтр is_done и сортировка sort/order как у пользователя."),
        ("Админ: чтение одной", "GET /admin/tasks/{id} — любая задача; в ответе owner_email владельца."),
        ("Админ: создание", "POST /admin/tasks — поля задачи + owner_id; 404 если пользователь с таким id не найден."),
        ("Админ: обновление", "PATCH /admin/tasks/{id} — частичное изменение любой задачи (схема TaskUpdate)."),
        ("Админ: удаление", "DELETE /admin/tasks/{id} — удаление любой задачи; только при роли admin."),
    ]
    for a, b in rows:
        row = tbl.add_row().cells
        row[0].text = a
        row[1].text = b
    _set_table_body_font(tbl)
    doc.add_paragraph()

    h = doc.add_heading("4. Описание стека технологий", level=1)
    _style_heading_runs(h, 16)
    _p(
        doc,
        "Ниже приведён стек реализации проекта (каталоги репозитория: backend/, web_client/, mobile/todo_app/).",
    )
    _bullets(
        doc,
        [
            "Backend: Python 3.11+, веб-фреймворк FastAPI, ASGI-сервер Uvicorn, валидация данных Pydantic v2, асинхронный слой SQLAlchemy 2.x (драйверы aiosqlite для SQLite и asyncpg для PostgreSQL).",
            "База данных: SQLite для локальной разработки (файл todo.db); PostgreSQL — опционально через docker-compose; схема версионируется Alembic.",
            "Безопасность: хэширование пароля (bcrypt), подпись и проверка JWT (python-jose).",
            "Наблюдаемость: эндпоинты /health и /metrics (Prometheus); опционально внешняя телеметрия Logfire при наличии токена в .env.",
            "Веб-клиент: одностраничное приложение на HTML5, CSS, JavaScript (fetch API), хранение JWT в localStorage.",
            "Мобильный и десктопный клиент: Flutter 3.x, пакеты http, flutter_secure_storage, intl; интерфейс согласован с веб-клиентом.",
            "Тестирование backend: pytest, HTTP-клиент FastAPI TestClient (интеграционные сценарии).",
            "Теоретическая база по FastAPI — в том числе материалы курса Stepik «FastAPI, Docker и CI/CD» (ссылка в корневом README проекта).",
        ],
    )

    h = doc.add_heading("5. Операции CRUD", level=1)
    _style_heading_runs(h, 16)
    _p(
        doc,
        "CRUD (Create, Read, Update, Delete) — канонический набор операций жизненного цикла над **одной сущностью** "
        "(здесь — «задача») в приложении. Формально: **C** — внесение новой записи в множество; **R** — извлечение "
        "одного элемента или подмножества без изменения состояния хранилища; **U** — изменение атрибутов существующей "
        "записи; **D** — исключение записи из множества. В архитектуре «тонкий контроллер — слой сервиса/репозитория — "
        "СУБД» эти операции отображаются на SQL-команды INSERT, SELECT, UPDATE, DELETE и на HTTP-методы REST.",
    )
    _p(
        doc,
        "В учебных материалах по Java часто рассматривают реализацию CRUD через Spring и REST-контроллеры "
        "(см., например, материалы JavaRush по Spring), с точки зрения пользователя — сценарии взаимодействия "
        "с формами и списками (обзор на Habr), а на уровне персистентности — через Hibernate и SQL "
        "(Skillbox). В данном проекте используется тот же логический цикл, но стек Python/FastAPI/SQLAlchemy.",
    )

    _h2(doc, "5.1. С точки зрения клиента (интерфейс и REST)")
    _p(
        doc,
        "Клиентом выступает браузер (веб-страница web_client/index.html, вызовы fetch) или приложение на Flutter "
        "(пакет http, класс TodoApi в lib/api.dart). Оба клиента работают с одним и тем же API: базовый URL задаётся "
        "пользователем (поле «Адрес API» / defaultBaseUrl), далее все защищённые запросы отправляются с заголовком "
        "Authorization: Bearer <access_token>, где токен получен после POST /auth/login.",
    )
    _p(doc, "Подготовка сессии и выбор маршрутов задач:", bold=True)
    _bullets(
        doc,
        [
            "Регистрация: POST /auth/register, тело JSON {email, password}; при успехе 201, при занятом email — 400.",
            "Вход: POST /auth/login с телом application/x-www-form-urlencoded (поле username = email, password); "
            "ответ 200 и JSON {access_token, token_type: bearer}.",
            "Профиль: GET /auth/me с тем же Bearer — ответ {id, email, role} из базы. Веб сохраняет роль и id в "
            "localStorage; Flutter — в secure storage. По role=admin клиенты переключают префикс задач с /tasks на "
            "/admin/tasks для списка, создания, PATCH и DELETE (полный CRUD по всем задачам).",
        ],
    )
    _p(doc, "Операции CRUD над задачей с позиции клиента (обычный пользователь, префикс /tasks):", bold=True)
    _bullets(
        doc,
        [
            "Create — POST /tasks, Content-Type: application/json, тело {title, description|null, is_done}. Успех 201, "
            "тело TaskRead (сервер подставляет owner_id текущего пользователя). Ошибки: 401 без токена, 422 при "
            "нарушении ограничений полей.",
            "Read (список) — GET /tasks с необязательными query: is_done (true|false), sort (created_at|status), "
            "order (asc|desc). Успех 200, массив TaskRead; у обычного пользователя поле owner_email в JSON обычно null.",
            "Read (одна) — GET /tasks/{id}. Успех 200 и одна TaskRead; 404 если задачи нет или она не принадлежит "
            "текущему пользователю (клиент не различает причину).",
            "Update — PATCH /tasks/{id}, тело только изменяемых полей (title, description, is_done — любое подмножество). "
            "Успех 200 и полный TaskRead после обновления.",
            "Delete — DELETE /tasks/{id}. Успех 204 без тела; 404 при отсутствии доступа; 401 при просроченном токене.",
        ],
    )
    _p(doc, "Те же операции для роли admin (префикс /admin/tasks):", bold=True)
    _bullets(
        doc,
        [
            "Read (список) — GET /admin/tasks с теми же query is_done, sort, order; ответ 200, в каждом элементе "
            "заполнены owner_id и owner_email владельца.",
            "Read (одна) — GET /admin/tasks/{id} по любому идентификатору задачи в системе.",
            "Create — POST /admin/tasks: к полям задачи добавляется owner_id (существующий пользователь); 404 если "
            "такого пользователя нет. В веб-клиенте при пустом поле владельца подставляется id из /auth/me; в Flutter — "
            "аналогично через сохранённый user id.",
            "Update и Delete — PATCH и DELETE /admin/tasks/{id} для любой задачи; при отсутствии роли admin — 403.",
        ],
    )
    _p(doc, "Поведение интерфейса (сопоставление с HTTP):", bold=True)
    _bullets(
        doc,
        [
            "Веб: после входа вызывается /auth/me, затем загрузка списка; фильтры и сортировки меняют только query к "
            "GET; чекбокс «выполнено» и inline-редактирование заголовка вызывают PATCH; кнопка удаления — DELETE; "
            "кнопка «Добавить» — POST. При 401 клиент очищает токен и показывает форму входа.",
            "Flutter: при старте с сохранённым токеном выполняется fetchMe, затем listTasks; переключение чекбокса, "
            "переименование и удаление вызывают patchTask/deleteTask с тем же выбором URL, что и в веб-клиенте.",
        ],
    )

    _h2(doc, "5.2. С точки зрения базы данных")
    _p(
        doc,
        "Реляционная модель: таблица users (идентификатор, уникальный email, хэш пароля, роль, метки времени); "
        "таблица tasks (идентификатор, заголовок, описание, признак выполнения, внешний ключ owner_id на users, "
        "временные метки создания и обновления). Операции INSERT/SELECT/UPDATE/DELETE над строками tasks выполняются "
        "через ORM с обязательным ограничением по owner_id для обычного пользователя.",
    )

    _h2(doc, "5.3. С точки зрения backend (стек FastAPI)")
    _p(
        doc,
        "Маршруты объявлены в модулях app/routers (например, tasks.py, auth.py, admin.py). Слой доступа к данным — "
        "app/crud.py: функции с явными именами (create_task, list_tasks_for_user, update_task и т.д.), что один в один "
        "отражает операции CRUD над таблицей tasks. Входящие и исходящие структуры описаны схемами Pydantic "
        "(app/schemas.py: TaskCreate, TaskUpdate, TaskRead). Перед выполнением операций над задачей из JWT "
        "извлекается текущий пользователь (app/deps.py); для обычных маршрутов /tasks/* доступ к чужой записи "
        "исключается на уровне запросов с условием owner_id.",
    )

    _h2(doc, "5.4. С точки зрения абстракции CRUD (ресурс, идемпотентность, границы)")
    _p(
        doc,
        "С позиции проектирования API **ресурс «задача»** идентифицируется целочисленным **id** (суррогатный первичный "
        "ключ в таблице tasks). Коллекция всех задач системы не имеет единого URL для обычного пользователя: "
        "пользователь работает только со своим подмножеством, задаваемым неявно через JWT.",
    )
    _p(doc, "Соответствие буквам CRUD и HTTP в проекте:", bold=True)
    _bullets(
        doc,
        [
            "**Create** — POST /tasks или POST /admin/tasks: тело описывает начальное состояние сущности "
            "(заголовок, описание, признак выполнения). Операция **не идемпотентна**: повторные запросы создают "
            "несколько разных задач. Идентификатор и метки времени назначает сервер. Владелец: для POST /tasks "
            "подставляется id из JWT; для POST /admin/tasks передаётся явный owner_id.",
            "**Read** — GET: либо коллекция (список), либо один элемент по id. Для списка поддерживаются параметры "
            "запроса is_done, sort, order — это часть операции чтения (фильтрация и упорядочивание представления "
            "коллекции). Ответ 200 с телом JSON. Отсутствие записи или права — 404 (для пользователя чужая задача "
            "неотличима от отсутствующей).",
            "**Update** — PATCH /tasks/{id} или PATCH /admin/tasks/{id}: **частичная** семантика (в теле только "
            "изменяемые поля согласно TaskUpdate). На уровне ORM обновляются только переданные атрибуты; ответ — "
            "полное текущее состояние TaskRead. Операция идемпотентна по смыслу только для фиксированного тела; "
            "сервер обновляет updated_at.",
            "**Delete** — DELETE: удаление строки в БД. Успех — **204 No Content** без тела. Повторный DELETE по "
            "тому же id после успешного удаления даёт 404. Для пользователя удалять можно только свои задачи.",
            "**Две поверхности REST** для одной сущности: префикс /tasks/* (изоляция по владельцу) и /admin/tasks/* "
            "(полный CRUD по всем задачам при роли admin). Это разграничение **авторизации**, а не разных моделей "
            "данных: таблица tasks одна, меняется набор допустимых строк и полей ответа (например, owner_email в "
            "админских ответах).",
        ],
    )
    _p(
        doc,
        "Таким образом, с точки зрения CRUD проект демонстрирует полный цикл над сущностью «задача» с разделением "
        "прав доступа между ролями user и admin без дублирования схемы БД.",
    )

    h = doc.add_heading("6. Выбор и использование линтера", level=1)
    _style_heading_runs(h, 16)
    _p(
        doc,
        "Линтер — инструмент статического анализа кода без его выполнения: он помогает находить ошибки, несоответствия "
        "стилю и потенциальные дефекты на ранней стадии (см. статью Яндекс Практикума о линтерах).",
    )
    _bullets(
        doc,
        [
            "Python: Ruff — объединяет проверки в духе flake8, isort и др.; правила заданы в backend/pyproject.toml (наборы E, F, I, UP, B). Запуск: py -3 -m ruff check app tests.",
            "Dart/Flutter: пакет flutter_lints и файл mobile/todo_app/analysis_options.yaml задают рекомендуемый стиль для клиентского кода.",
        ],
    )

    h = doc.add_heading("7. Выбор и использование СУБД", level=1)
    _style_heading_runs(h, 16)
    _p(
        doc,
        "В формулировках учебных заданий иногда используется аббревиатура «СУЗ»; в контексте хранения данных "
        "подразумевается система управления базами данных (СУБД).",
    )
    _bullets(
        doc,
        [
            "Для разработки и сдачи без отдельного сервера БД выбран SQLite: один файл базы, простое подключение через строку DATABASE_URL в .env.",
            "Для сценария, близкого к промышленному, предусмотрен PostgreSQL (docker-compose в корне репозитория, драйвер asyncpg, миграции Alembic).",
            "Обоснование: SQLite минимизирует порог входа; PostgreSQL подходит при конкурентной нагрузке и сложных запросах.",
        ],
    )

    h = doc.add_heading(
        "8. Системы управления знаниями и инструменты документирования проекта",
        level=1,
    )
    _style_heading_runs(h, 16)
    _p(
        doc,
        "В отчётных формулировках под системами управления знаниями (СУЗ) в ИТ часто понимают средства фиксации и "
        "совместного использования знаний о продукте и процессе разработки; это не следует смешивать с СУБД из раздела 7, "
        "где речь идёт о хранении данных приложения.",
    )
    _p(doc, "Примеры класса организационных СУЗ (для отчёта, при необходимости):", bold=True)
    _bullets(
        doc,
        [
            "Корпоративные вики и порталы: Confluence, Notion, Microsoft SharePoint; при командной работе — также Wiki в GitHub.",
            "В данном учебном проекте основной акцент — репозиторий с исходным кодом и автоматически поддерживаемое описание REST API.",
        ],
    )
    _p(doc, "Инженерные носители знаний, используемые в проекте:", bold=True)
    _bullets(
        doc,
        [
            "Git и GitHub — версионирование кода, история изменений, публикация репозитория; сопутствующие знания — в README.md и каталоге docs/.",
            "OpenAPI — машиночитаемое описание HTTP API. FastAPI формирует спецификацию (эндпоинт /openapi.json) и интерактивные страницы /docs (Swagger UI) и /redoc (ReDoc), что согласует документацию с фактическим поведением сервера и облегчает сопровождение веб- и Flutter-клиентов.",
        ],
    )

    h = doc.add_heading("9. Анализ кода на соответствие принципам SOLID", level=1)
    _style_heading_runs(h, 16)
    _p(
        doc,
        "Принципы SOLID (единственная ответственность, открытость/закрытость, подстановка Лисков, разделение интерфейса, "
        "инверсия зависимостей) применимы к объектно-ориентированному коду на любом языке; разбор на примере JavaScript "
        "можно использовать как методологию (Sprintcode). Ниже — краткая проекция на слои backend-проекта.",
    )
    _bullets(
        doc,
        [
            "SRP: роутеры отвечают за HTTP-контракт; crud.py — за операции с БД; security.py — за хэширование и JWT; schemas — за формы данных.",
            "OCP: добавление новых эндпоинтов не требует изменения низкоуровневого кода драйвера БД — расширение через новые функции и маршруты.",
            "LSP: схемы Pydantic и ORM-модели согласованы через model_validate и согласованные поля.",
            "ISP: зависимости эндпоинтов узкие (сессия БД, текущий пользователь), без избыточных «толстых» интерфейсов.",
            "DIP: приложение зависит от абстракций get_db и набора функций слоя crud, а не от деталей конкретного драйвера.",
        ],
    )

    h = doc.add_heading("10. Метод авторизации на backend и подходы к контролю доступа", level=1)
    _style_heading_runs(h, 16)
    _p(
        doc,
        "В проекте реализованы аутентификация по паре email/пароль с последующей выдачей JWT и авторизация на основе "
        "ролей (RBAC): обычный пользователь и администратор.",
    )
    _p(doc, "Подходы к контролю доступа в теории:", bold=True)
    _bullets(
        doc,
        [
            "RBAC (Role-Based Access Control): права привязаны к ролям субъекта; проверка «имеет ли пользователь роль admin» для административных маршрутов.",
            "ABAC (Attribute-Based Access Control): решения на основе атрибутов субъекта, объекта и контекста; в учебном TODO не используется из-за избыточности.",
            "Обзоры RBAC vs ABAC: материалы Glabit, Keeper Security; теория также представлена на Stepik (уроки курса по смежной тематике).",
        ],
    )
    _p(doc, "Практические библиотеки (ориентиры для промышленных систем):", bold=True)
    _bullets(
        doc,
        [
            "Casbin — универсальный движок политик доступа (статьи на Habr).",
            "RBACX для Python — пример лёгкой реализации RBAC (статья на Habr). В данном проекте достаточно проверки роли из JWT и зависимости require_admin без отдельной библиотеки политик.",
        ],
    )

    h = doc.add_heading("11. Выбор и использование логгера", level=1)
    _style_heading_runs(h, 16)
    _p(
        doc,
        "Логирование — фиксация событий работы приложения для отладки, аудита и эксплуатации. В Python используется "
        "стандартный модуль logging; конфигурация может выводить структурированные JSON-сообщения в stdout.",
    )
    _bullets(
        doc,
        [
            "Теория назначения логгера и уровней логирования — материал KursHub.",
            "Обзор инструментов и практик — публикации на Habr (UltravDS, МТС и др.).",
            "В проекте логирование применяется при обработке запросов и значимых операциях (создание/изменение задач, списки для администратора) — см. модуль настройки логирования в backend.",
        ],
    )

    h = doc.add_heading("12. Тестирование", level=1)
    _style_heading_runs(h, 16)
    _p(
        doc,
        "На frontend в экосистеме React распространены тесты компонентов и интеграционные сценарии (например, материалы "
        "PurpleSchool по React Testing Library и смежным инструментам). В проекте на Flutter используется пакет flutter_test "
        "для базовой проверки монтирования приложения; для backend — pytest и TestClient FastAPI: проверяются здоровье сервиса, "
        "полный цикл CRUD задач и ограничение доступа к административному API для роли user.",
    )

    h = doc.add_heading("13. Репозиторий проекта и установка", level=1)
    _style_heading_runs(h, 16)
    _p(
        doc,
        "Исходный код учебного проекта следует разместить в открытом репозитории на GitHub; ниже приведена инструкция по "
        "локальной установке и запуску (после клонирования репозитория).",
    )
    _p(doc, "Клонирование и общая структура:", bold=True)
    _bullets(
        doc,
        [
            "git clone <URL-репозитория> и переход в каталог проекта.",
            "Backend: cd backend; создание виртуального окружения; pip install -r requirements.txt; копирование .env.example в .env; настройка SECRET_KEY и при необходимости DATABASE_URL.",
            "Запуск API: py -3 -m uvicorn app.main:app --reload --host 0.0.0.0 --port 8000.",
            "Веб-клиент: открыть web_client/index.html в браузере или поднять статический сервер (например, python -m http.server в каталоге web_client).",
            "Flutter: cd mobile/todo_app; flutter pub get; flutter run (или flutter run -d chrome для веба).",
        ],
    )
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(1.25)
    rph = p.add_run(
        "Актуальный URL репозитория на GitHub: указать после публикации (заменить placeholder в документе)."
    )
    _set_run_font(rph, size_pt=14, color=RGB_BLACK)

    h = doc.add_heading("14. Список источников", level=1)
    _style_heading_runs(h, 16)
    sources = [
        ("Курс Stepik «FastAPI, Docker и CI/CD» (промо)", "https://stepik.org/course/179694/promo"),
        ("Пример CRUD на Java Spring (JavaRush)", "https://javarush.com/quests/lectures/ru.javarush.java.spring.lecture.level08.lecture05"),
        ("CRUD с точки зрения пользователя (Habr)", "https://habr.com/ru/articles/971138/"),
        ("CRUD на Hibernate для начинающих (Skillbox)", "https://skillbox.ru/media/code/crudprilozhenie-na-hibernate-dlya-nachinayushchikh/"),
        ("Что такое линтер (Яндекс Практикум)", "https://practicum.yandex.ru/blog/chto-takoe-linter-v-programmirovanii/"),
        ("SOLID на примере JS (Sprintcode)", "https://sprintcode.pro/ru/blog/solid-principles"),
        ("Подходы к контролю доступа RBAC vs ABAC (Glabit)", "https://glabit.ru/blog/podhody-k-kontrolyu-dostupa-rbac-vs-abac"),
        ("RBAC vs ABAC — обзор (Keeper Security)", "https://www.keepersecurity.com/blog/ru/2024/10/28/rbac-vs-abac-which-should-you-use/"),
        ("Stepik: урок по смежной тематике", "https://stepik.org/lesson/2045912/step/1?unit=2074376"),
        ("Библиотека Casbin (Habr)", "https://habr.com/ru/articles/823374/"),
        ("RBACX для Python (Habr)", "https://habr.com/ru/articles/950080/"),
        ("Что такое logger (KursHub)", "https://kurshub.ru/journal/blog/chto-takoe-logger-v-programmirovanii-i-tehnike-polnoe-obyasnenie-prostymi-slovami/"),
        ("Обзор инструментов логирования (Habr, UltravDS)", "https://habr.com/ru/companies/ultravds/articles/984372/"),
        ("Пример и практики логирования (Habr, МТС)", "https://habr.com/ru/companies/ru_mts/articles/715212/"),
        ("Тестирование React (PurpleSchool)", "https://purpleschool.ru/knowledge-base/react/dev/react-test"),
        ("Документация Flutter", "https://flutter.dev/"),
        ("Документация FastAPI", "https://fastapi.tiangolo.com/"),
        ("Документация Ruff", "https://docs.astral.sh/ruff/"),
        ("Спецификация OpenAPI", "https://spec.openapis.org/oas/latest.html"),
        ("Автоматическая документация API в FastAPI (Swagger UI / ReDoc)", "https://fastapi.tiangolo.com/features/#automatic-docs"),
        ("GitHub", "https://github.com/"),
    ]
    for i, (title, url) in enumerate(sources, 1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.add_run(f"{title}. ")
        _add_hyperlink(p, url, url)
        for run in p.runs:
            _set_run_font(run, size_pt=14, color=RGB_BLACK)

    _apply_page_margins(doc)
    return doc


def main() -> None:
    out = Path(__file__).resolve().parent / "POYASNITELNAYA_ZAPISKA.docx"
    doc = build()
    doc.save(out)
    print(f"Written: {out}")


if __name__ == "__main__":
    main()
